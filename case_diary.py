# FULL WORD AUTOMATION SYSTEM (Python) — FIXED
# Generates a Case Diary in .docx if python-docx is available,
# otherwise falls back to a Word-compatible .rtf file (opens in MS Word).

from datetime import datetime

# Plain-text case diary template (inserted per request). This version uses
# named placeholders so fields can be filled while preserving appearance.
CASE_DIARY_TEMPLATE = """Schedule XLVII-Form No.120
P.M. Form No.80

Case Diary
(Rule 104)

Police Station- {police_station:25}    District- {district}

First information Report No.- {fir_no:20}    CD No.- {cd_no:8}    U/s- {sections}

Date (with hour) on which action was taken and places visited.
Dt.- {date:10}   at {time}

Opening at {ps_name:20} PS	Record of Investigation
    PO  {po}
    DO  {do}
    DR  {dr}
    Accd  {accd}
    I.O  {io}
    Opened the diary for the day and commenced investigation of the case.

Gist of FIR
{gist}

Registration of Case		{registration}

Examination of Complt.
        Examined him. He is the Complt. of this case. His further statement has been recorded U/S 180(3) BNSS in a separate sheet which is enclosed herewith.
    Left for Spot i.e. {left_time}
    Reached at spot {reach_time}

Spot visit		Visited the spot of the case being identified by the complainant. The spot of the case is
{spot}

Examination of P/W
{witness1}

Examination of P/W
{witness2}

Search for other witnesses
{search}

Discussion with IIC
    Discussed with my IIC regarding the progress of the investigation and steps taken by me till date in this case.

Close		{close}

                                      Submitted


                    (the name and
                      designation of the IO
                      Police Station’s name)
"""


def fill_template(data: dict) -> str:
    """Return the CASE_DIARY_TEMPLATE filled with values from `data`.

    Missing keys are rendered as empty strings.
    """
    tpl_values = {k: data.get(k, '') for k in (
        'police_station', 'district', 'fir_no', 'cd_no', 'sections',
        'date', 'time', 'ps_name', 'po', 'do', 'dr', 'accd', 'io',
        'gist', 'registration', 'complainant', 'left_time', 'reach_time',
        'spot', 'witness1', 'witness2', 'search', 'discussion', 'close',
        'io_name', 'designation', 'ps'
    )}
    return CASE_DIARY_TEMPLATE.format(**tpl_values)

# ------------------------------
# OPTIONAL IMPORT (with fallback)
# ------------------------------
try:
    from docx import Document  # python-docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# ------------------------------
# VALIDATION
# ------------------------------
REQUIRED_KEYS = [
    "police_station", "district", "fir_no", "cd_no", "sections",
    "date", "time", "ps_name", "po", "do", "dr", "accd", "io",
    "gist", "registration", "complainant", "left_time", "reach_time",
    "spot", "witness1", "witness2", "search", "discussion", "close",
    "io_name", "designation", "ps"
]


def validate_input(data: dict):
    missing = [k for k in REQUIRED_KEYS if k not in data]
    if missing:
        raise ValueError(f"Missing required fields: {', '.join(missing)}")


# ------------------------------
# CORE RENDER (shared text)
# ------------------------------

def build_lines(data: dict):
    lines = []
    def add(t=""):
        lines.append(t)

    add("Schedule XLVII-Form No.120")
    add("P.M. Form No.80")
    add("")
    add("Case Diary (Rule 104)")
    add("")

    add(f"Police Station- {data['police_station']}")
    add(f"District- {data['district']}")
    add("")

    add(f"FIR No.- {data['fir_no']}    CD No.- {data['cd_no']}    U/s- {data['sections']}")
    add("")

    add(f"Date: {data['date']}   Time: {data['time']}")
    add("")

    add(f"Opening at {data['ps_name']}")
    add("Record of Investigation")
    add("")

    add(f"PO: {data['po']}")
    add(f"DO: {data['do']}")
    add(f"DR: {data['dr']}")
    add(f"Accd: {data['accd']}")
    add(f"I.O: {data['io']}")
    add("")

    add("Opened the diary for the day and commenced investigation.")
    add("")

    add("Gist of FIR:")
    add(data['gist'])
    add("")

    add("Registration of Case:")
    add(data['registration'])
    add("")

    add("Examination of Complainant:")
    add(data['complainant'])
    add("")

    add(f"Left for Spot: {data['left_time']}")
    add(f"Reached at Spot: {data['reach_time']}")
    add("")

    add("Spot Visit:")
    add(data['spot'])
    add("")

    add("Examination of P/W:")
    add(data['witness1'])
    add("")

    add("Examination of P/W:")
    add(data['witness2'])
    add("")

    add("Search for other witnesses:")
    add(data['search'])
    add("")

    add("Discussion with IIC:")
    add(data['discussion'])
    add("")

    add("Close:")
    add(data['close'])
    add("")

    add("Submitted")
    add("")

    add(data['io_name'])
    add(data['designation'])
    add(data['ps'])

    return lines


# ------------------------------
# DOCX GENERATION (if available)
# ------------------------------

def generate_docx(lines, output_file):
    # Backwards-compatible paragraph writer (kept for callers that pass lines).
    doc = Document()
    for ln in lines:
        doc.add_paragraph(ln)
    doc.save(output_file)


def generate_docx_from_data(data: dict, output_file: str):
    """Generate a .docx using a table layout that mirrors the provided template.

    This creates a headed document and a two-column table where left column
    contains labels and right column contains values (multi-line fields span
    the right cell).
    """
    doc = Document()

    # Use the plain-text template appearance and write each line as a paragraph
    filled = fill_template(data)
    for ln in filled.splitlines():
        p = doc.add_paragraph(ln)
    doc.save(output_file)


# ------------------------------
# RTF FALLBACK (no dependency)
# ------------------------------

def escape_rtf(text: str) -> str:
    return text.replace('\\', r'\\').replace('{', r'\\{').replace('}', r'\\}')


def generate_rtf(lines, output_file):
    # Basic RTF document. Word opens it natively.
    body = r"\\par ".join(escape_rtf(ln) for ln in lines)
    rtf = r"{\\rtf1\\ansi\\deff0 " + body + r"}"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(rtf)


def generate_rtf_from_data(data: dict, output_file: str):
    """Generate RTF using the textual template appearance."""
    filled = fill_template(data)
    lines = filled.splitlines()
    generate_rtf(lines, output_file)


# ------------------------------
# PUBLIC API
# ------------------------------

def generate_case_diary(data, output_file_base="case_diary_output"):
    """
    Generates:
    - output_file_base.docx if python-docx is available
    - else output_file_base.rtf
    Returns the output file path.
    """
    validate_input(data)

    if DOCX_AVAILABLE:
        out = output_file_base + ".docx"
        generate_docx_from_data(data, out)
    else:
        out = output_file_base + ".rtf"
        generate_rtf_from_data(data, out)

    print(f"Document generated: {out} (DOCX_AVAILABLE={DOCX_AVAILABLE})")
    return out


# ------------------------------
# SAMPLE INPUT
# ------------------------------

def sample_data_now():
    return {
        "police_station": "Badabazar PS",
        "district": "Ganjam",
        "fir_no": "123/2026",
        "cd_no": "1",
        "sections": "BNS 137(2)",
        "date": datetime.now().strftime("%d/%m/%Y"),
        "time": datetime.now().strftime("%I:%M %p"),
        "ps_name": "Badabazar PS",
        "po": "N/A",
        "do": "N/A",
        "dr": "N/A",
        "accd": "N/A",
        "io": "SI Chinmaya Sahu",
        "gist": "Brief facts of the FIR...",
        "registration": "Basing on report, case registered and investigation taken up.",
        "complainant": "Examined complainant; statement recorded U/S 180(3) BNSS.",
        "left_time": "10:30 AM",
        "reach_time": "11:15 AM",
        "spot": "Visited spot identified by complainant.",
        "witness1": "Witness 1 examined.",
        "witness2": "Witness 2 examined.",
        "search": "Search conducted for additional witnesses.",
        "discussion": "Discussed progress with IIC.",
        "close": "Closed the diary for the day pending further investigation.",
        "io_name": "SI Chinmaya Sahu",
        "designation": "Investigating Officer",
        "ps": "Badabazar Police Station"
    }


# ------------------------------
# TESTS
# ------------------------------

def _test_generation():
    data = sample_data_now()
    out = generate_case_diary(data, output_file_base="test_case_diary")
    # Check file exists
    import os
    assert os.path.exists(out), "Output file was not created"
    # Basic content sanity (read back for RTF only)
    if out.endswith('.rtf'):
        with open(out, 'r', encoding='utf-8') as f:
            content = f.read()
        assert 'Case Diary (Rule 104)' in content
    print("All tests passed.")


if __name__ == "__main__":
    # Run sample generation
    data = sample_data_now()
    generate_case_diary(data)

    # Run tests
    _test_generation()


# ------------------------------
# HOW TO USE
# ------------------------------
# 1) If you want .docx output:
#    pip install python-docx
#
# 2) Run:
#    python script.py
#
# 3) If python-docx is not installed, the script will still work
#    and produce a .rtf file that opens in MS Word.

# ------------------------------
# NOTES / NEXT STEPS
# ------------------------------
# - Replace sample_data_now() with form input (Flask/FastAPI)
# - Add digital signature / PDF export
# - Add audit trail fields (edited_by, timestamps)
# - Lock entries after generation for legal integrity
